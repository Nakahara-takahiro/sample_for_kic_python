import docx
doc = docx.Document('sample.docx')
num = 0

para2 = doc.paragraphs[0]
t = para2.text
t = t.replace('ダミー','エラー')
para2.text = t

'''
for para in doc.paragraphs:
    num = num + 1
    print(num)
    para2 = doc.paragraphs[num]
    t = para2.text
    t = t.replace('ダミー','エラー')
    para2.text = t
'''
doc.save('sample2.docx')